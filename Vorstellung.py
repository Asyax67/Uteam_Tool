# Vorstellung.py

import os
from datetime import datetime
from docx import Document
from openpyxl import load_workbook

from PyQt6.QtWidgets import (
    QWidget, QLabel, QLineEdit, QTextEdit, QComboBox, QPushButton,
    QMessageBox, QGridLayout, QHBoxLayout, QVBoxLayout,
    QSizePolicy, QFrame, QScrollArea
)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QFont

# ---- Passe hier deine Pfade an: ----
EXCEL_PATH = r"C:\Users\ASYAKKA\Mercedes-Benz (corpdir.onmicrosoft.com)\DWT_UTeam Werk 10 - General\08_Rotation UTeam\Projekt_UTeam_Digitalisierung\Masterliste_UTeam.xlsx"
import os

# Statt des OneDrive‑Pfads:
# OUTPUT_DIR = r"C:\Users\ASYAKKA\...\Vorstellung_Asya_Test"
# Nimm stattdessen:
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "Vorstellung_Asya_Test")

# -------------------------------------

class Vorstellung(QWidget):
    def __init__(self, main_window):
        super().__init__()
        self.main_window = main_window

        # Scroll-Area als Haupt-Widget
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)

        # Container für alle Formular-Widgets
        container = QWidget()
        scroll.setWidget(container)

        # Main-Layout dieser Seite
        outer = QVBoxLayout(self)
        outer.addWidget(scroll)

        # Inneres Layout auf dem Container
        main_layout = QVBoxLayout(container)
        main_layout.setContentsMargins(30,30,30,30)
        main_layout.setSpacing(20)

        # = Header: Titel + Datum =
        header = QHBoxLayout()
        header.setAlignment(Qt.AlignmentFlag.AlignVCenter)

        title = QLabel("Vorstellungsgespräch")
        title.setFont(QFont("Helvetica", 20, QFont.Weight.Bold))
        header.addWidget(title)
        header.addStretch()

        header.addWidget(QLabel("Datum:"))
        self.le_date = QLineEdit()
        self.le_date.setPlaceholderText("DD.MM.YYYY")
        self.le_date.setMaximumWidth(120)
        header.addWidget(self.le_date)

        main_layout.addLayout(header)

        # Trennlinie
        sep = QFrame(); sep.setFrameShape(QFrame.Shape.HLine)
        main_layout.addWidget(sep)

        # = Grid mit Pflichtfeldern und Stammdaten =
        grid = QGridLayout()
        grid.setHorizontalSpacing(20)
        grid.setVerticalSpacing(10)
        main_layout.addLayout(grid)

        # Zeile 0: Vorname | Nachname
        self.le_vor  = QLineEdit()
        self.le_nach = QLineEdit()
        for w in (self.le_vor, self.le_nach):
            w.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        grid.addWidget(QLabel("Vorname *"),  0,0, alignment=Qt.AlignmentFlag.AlignRight)
        grid.addWidget(self.le_vor,           0,1)
        grid.addWidget(QLabel("Nachname *"),  0,2, alignment=Qt.AlignmentFlag.AlignRight)
        grid.addWidget(self.le_nach,          0,3)

        # Zeile 1: Geburtsdatum | Wohnort
        self.le_geb  = QLineEdit(); self.le_geb.setPlaceholderText("DD.MM.YYYY")
        self.le_wohn = QLineEdit()
        for w in (self.le_geb, self.le_wohn):
            w.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        grid.addWidget(QLabel("Geburtsdatum *"),1,0, alignment=Qt.AlignmentFlag.AlignRight)
        grid.addWidget(self.le_geb,             1,1)
        grid.addWidget(QLabel("Wohnort"),       1,2, alignment=Qt.AlignmentFlag.AlignRight)
        grid.addWidget(self.le_wohn,            1,3)

        # Zeile 2: Aktueller Einsatz | Stamm-Kostenstelle
        self.le_eins = QLineEdit()
        self.le_kst  = QLineEdit()
        for w in (self.le_eins, self.le_kst):
            w.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        grid.addWidget(QLabel("Aktueller Einsatz *"),   2,0, alignment=Qt.AlignmentFlag.AlignRight)
        grid.addWidget(self.le_eins,                    2,1)
        grid.addWidget(QLabel("Stamm-Kostenstelle *"),  2,2, alignment=Qt.AlignmentFlag.AlignRight)
        grid.addWidget(self.le_kst,                     2,3)

        # Zeile 3: Geschlecht
        self.cb_gender = QComboBox()
        self.cb_gender.addItems(["m","w","d"])
        grid.addWidget(QLabel("Geschlecht"), 3,0, alignment=Qt.AlignmentFlag.AlignRight)
        grid.addWidget(self.cb_gender,       3,1)

        # Zeile 4: Staplerschein
        self.cb_fork = QComboBox()
        self.cb_fork.addItems(["ja","nein","k.A."])
        grid.addWidget(QLabel("Staplerschein"), 4,0, alignment=Qt.AlignmentFlag.AlignRight)
        grid.addWidget(self.cb_fork,            4,1)

        # Spalten 1 und 3 dehnen
        grid.setColumnStretch(1,1)
        grid.setColumnStretch(3,1)

        # Trennlinie
        sep2 = QFrame(); sep2.setFrameShape(QFrame.Shape.HLine)
        main_layout.addWidget(sep2)

        # = große Textblöcke =
        for label, widget in [
            ("Laufbahn",      QTextEdit()),
            ("Qualifikation", QTextEdit()),
            ("Wunsch",        QTextEdit()),
            ("Sonstiges",     QTextEdit()),
        ]:
            widget.setFixedHeight(100)
            widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            setattr(self, f"te_{label.lower()}", widget)
            main_layout.addWidget(QLabel(label))
            main_layout.addWidget(widget)

        # Pflichtfeld-Hinweis
        lbl_pf = QLabel("* Pflichtfelder"); lbl_pf.setStyleSheet("color: darkred;")
        main_layout.addWidget(lbl_pf)

        # = Button-Leiste =
        btns = QHBoxLayout(); btns.addStretch()
        btn_save = QPushButton("Gespräch anlegen")
        btn_save.setStyleSheet("background-color:#4CAF50;color:white;padding:8px 20px;")
        btn_save.clicked.connect(self.save_vorstellung)
        btns.addWidget(btn_save)

        btn_clear = QPushButton("Formular leeren")
        btn_clear.setStyleSheet("background-color:#FFC107;color:white;padding:8px 20px;")
        btn_clear.clicked.connect(self.clear_form)
        btns.addWidget(btn_clear)

        btn_cancel = QPushButton("Abbrechen")
        btn_cancel.setStyleSheet("background-color:#9E9E9E;color:white;padding:8px 20px;")
        btn_cancel.clicked.connect(self.on_cancel)
        btns.addWidget(btn_cancel)

        main_layout.addLayout(btns)

    def clear_form(self):
        """Alle Eingabefelder zurücksetzen."""
        for w in (
            self.le_date, self.le_vor,  self.le_nach,
            self.le_geb,  self.le_wohn,
            self.le_eins, self.le_kst,
            self.te_laufbahn, self.te_qualifikation,
            self.te_wunsch,    self.te_sonstiges
        ):
            w.clear()
        self.cb_gender.setCurrentIndex(0)
        self.cb_fork.setCurrentIndex(0)

    def on_cancel(self):
        """Formular leeren und zurück auf Startseite."""
        self.clear_form()
        self.main_window.zeige_startseite()

    def save_vorstellung(self):
        """Word-Dokument erstellen und Excel ergänzen."""
        # 1) Pflicht prüfen
        if not all([
            self.le_vor.text().strip(),
            self.le_nach.text().strip(),
            self.le_geb.text().strip(),
            self.le_eins.text().strip(),
            self.le_kst.text().strip()
        ]):
            QMessageBox.warning(self, "Fehler", "Bitte alle Pflichtfelder ausfüllen.")
            return

        # 2) Werte sammeln
        date_str = self.le_date.text().strip() or datetime.today().strftime("%d.%m.%Y")
        vor    = self.le_vor.text().strip()
        nach   = self.le_nach.text().strip()
        geb    = self.le_geb.text().strip()
        eins   = self.le_eins.text().strip()
        kst    = self.le_kst.text().strip()
        gender = self.cb_gender.currentText()
        fork   = self.cb_fork.currentText()
        lauf   = self.te_laufbahn.toPlainText().strip()
        qual   = self.te_qualifikation.toPlainText().strip()
        wunsch = self.te_wunsch.toPlainText().strip()
        sonst  = self.te_sonstiges.toPlainText().strip()

        # 3) Word erzeugen
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        fname = f"{nach}_{vor}_{date_str.replace('.','-')}.docx"
        fullpath = os.path.join(OUTPUT_DIR, fname)
        try:
            doc = Document()
            doc.add_heading("Vorstellungsgespräch", level=1)
            doc.add_paragraph(f"Datum: {date_str}")
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Light List Accent 1"

            def add_row(lbl, val, bold=True):
                cells = tbl.add_row().cells
                run = cells[0].paragraphs[0].add_run(lbl)
                if bold:
                    run.bold = True
                cells[1].text = val

            # Beschriftungen + Werte
            for lbl, val, b in [
                ("Vorname:",        vor,    True),
                ("Nachname:",       nach,   True),
                ("Geburtsdatum:",   geb,    True),
                ("Aktueller Einsatz:", eins, True),
                ("Stamm-Kostenstelle:", kst,True),
                ("Geschlecht:",     gender, False),
                ("Staplerschein:",  fork,   False),
                ("Laufbahn:",       lauf,   False),
                ("Qualifikation:",  qual,   False),
                ("Wunsch:",         wunsch, False),
                ("Sonstiges:",      sonst,  False),
            ]:
                add_row(lbl, val, b)

            doc.save(fullpath)
        except Exception as e:
            QMessageBox.critical(self, "Word-Fehler", str(e))
            return

        # 4) Excel ergänzen (Spalten A–K)
        try:
            wb = load_workbook(EXCEL_PATH)
            ws = wb["Masterlist"]
            ws.append([
                vor, nach,      # A, B
                None,           # C (lfd. Nr.) → Excel-Formel oder manuell
                geb,            # D
                eins,           # E
                kst,            # F
                # falls du Stammcenter/Abteilung brauchst:
                None,           # G (Aktuelle Synopse)
                None,           # H (Stammcenter,-Abteilung)
                None,           # I (Stamm-Kostenstelle nochmals)
                gender,         # J
                fork            # K
            ])
            wb.save(EXCEL_PATH)
        except Exception as e:
            QMessageBox.critical(self, "Excel-Fehler", str(e))
            return

        # 5) Erfolg & zurück
        QMessageBox.information(self, "Erfolg",
            f"Word gespeichert unter:\n{fullpath}\n\nExcel-Eintrag hinzugefügt.")
        self.clear_form()
        self.main_window.zeige_startseite()
