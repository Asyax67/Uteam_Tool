# streamlit_app.py

import streamlit as st
import pandas as pd
from modul_rotation import lade_daten, finde_aktuellen_bereich
from openpyxl import load_workbook
from datetime import datetime
import os

# Pfad zur Masterliste
# statt absoluten Pfads
# EXCEL_PATH = r"C:\Users\…\Masterliste_UTeam.xlsx"

# benutze den relativen Pfad in Deinem Repo:
EXCEL_PATH = os.path.join(os.path.dirname(__file__), "Masterliste_UTeam.xlsx")

# Ordner, in den die Word‑Dokumente abgelegt werden
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "Vorstellung_Asya_Test")


# ─── Sidebar & Navigation ─────────────────────────────────────────────────────

st.sidebar.title("UTeam Rotations‑Tool")
page = st.sidebar.radio(
    "Navigation",
    ["🏠 Startseite", "👥 Übersicht", "🔄 Rotationsplan", "📝 Vorstellungsgespräch", "➕ Neuen Bereich anlegen"]
)


# ─── 1) Startseite ─────────────────────────────────────────────────────────────

if page == "🏠 Startseite":
    st.title("UTeam Rotations‑Tool")
    st.write("Wähle links einen Bereich aus, um fortzufahren.")


# ─── 2) Mitarbeiterübersicht ───────────────────────────────────────────────────

elif page == "👥 Übersicht":
    st.title("Mitarbeiterübersicht")
    df = lade_daten()
    if df is None:
        st.error("Excel konnte nicht geladen werden.")
    else:
        df["Aktueller Bereich"] = df.apply(finde_aktuellen_bereich, axis=1)
        df["Aktuelles Austrittsdatum"] = (
            pd.to_datetime(df["Aktuelles Austrittsdatum"], errors="coerce")
              .dt.strftime("%d.%m.%Y")
        )
        st.dataframe(df[["Vorname","Nachname","Aktueller Bereich","Aktuelles Austrittsdatum"]], height=500)


# ─── 3) Rotationsplan ───────────────────────────────────────────────────────────

elif page == "🔄 Rotationsplan":
    st.title("Rotationsplan")
    df = lade_daten(sheet_name="Rotationsplan")
    if df is None:
        st.error("Rotationsplan‑Tab konnte nicht geladen werden.")
    else:
        st.dataframe(df, height=500)


# ─── 4) Vorstellungsgespräch anlegen ───────────────────────────────────────────

elif page == "📝 Vorstellungsgespräch":
    st.title("Vorstellungsgespräch anlegen")

    with st.form("form_vorstellung"):
        col1, col2 = st.columns(2)
        with col1:
            vor   = st.text_input("Vorname *")
            nach  = st.text_input("Nachname *")
            geb   = st.text_input("Geburtsdatum * (DD.MM.YYYY)")
            eins  = st.text_input("Aktueller Einsatz *")
            kst   = st.text_input("Stamm‑Kostenstelle *", value="010-")
            gender = st.selectbox("Geschlecht", ["m","w","d"])
            fork   = st.selectbox("Staplerschein", ["ja","nein","k.A."])
        with col2:
            lauf   = st.text_area("Laufbahn")
            qual   = st.text_area("Qualifikation")
            wunsch = st.text_area("Wunsch")
            sonst  = st.text_area("Sonstiges")

        submitted = st.form_submit_button("Gespräch anlegen")
        if submitted:
            # --- Word erzeugen ---
            os.makedirs(OUTPUT_DIR, exist_ok=True)
            date_str = datetime.today().strftime("%d.%m.%Y")
            fname = f"{nach}_{vor}_{date_str.replace('.','-')}.docx"
            fullpath = os.path.join(OUTPUT_DIR, fname)

            from docx import Document
            doc = Document()
            doc.add_heading("Vorstellungsgespräch", level=1)
            doc.add_paragraph(f"Datum: {date_str}")

            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Light List Accent 1"
            def add_row(label, val, bold=True):
                cells = tbl.add_row().cells
                run = cells[0].paragraphs[0].add_run(label)
                if bold: run.bold = True
                cells[1].text = str(val)

            for lbl, val, b in [
                ("Vorname:",           vor,   True),
                ("Nachname:",          nach,  True),
                ("Geburtsdatum:",      geb,   True),
                ("Aktueller Einsatz:", eins,  True),
                ("Stamm‑Kostenstelle:",kst,   True),
                ("Geschlecht:",        gender,True),
                ("Staplerschein:",     fork,  True),
                ("Laufbahn:",          lauf,  False),
                ("Qualifikation:",     qual,  False),
                ("Wunsch:",            wunsch,False),
                ("Sonstiges:",         sonst, False),
            ]:
                add_row(lbl, val, b)

            doc.save(fullpath)

            # --- Excel ergänzen (Masterlist) ---
            wb = load_workbook(EXCEL_PATH)
            ws = wb["Masterlist"]
            # Spalten A=Vorname, B=Nachname, D=Geburtsdatum
            ws.append([vor, nach, None, geb])
            wb.save(EXCEL_PATH)

            st.success(f"Word: {fullpath}\nMasterlist aktualisiert.")


# ─── 5) Neuen Bereich anlegen ───────────────────────────────────────────────────

elif page == "➕ Neuen Bereich anlegen":
    st.title("Neuen Bereich anlegen")

    with st.form("form_bereich"):
        kategorie = st.selectbox(
            "Kategorie",
            ["eATS", "Gießerei", "Montage", "Logistik", "Qualität", "Fertigung"]
        )
        bereichsname = st.text_input("Name des Bereichs")
        kostenstelle = st.text_input("Kostenstelle", value="010-")

        save = st.form_submit_button("Bereich speichern")
        if save:
            try:
                wb = load_workbook(EXCEL_PATH)
                ws = wb["Rotationsplan"]  # hier wird in dein Rotationsplan-Blatt geschrieben

                # Spaltenindex ermitteln (1‑based)
                headers = [cell.value for cell in ws[1]]
                if kategorie not in headers:
                    st.error(f"Spalte '{kategorie}' nicht gefunden.")
                else:
                    col_idx = headers.index(kategorie) + 1

                    # erste freie Zelle in dieser Spalte finden
                    row_idx = 2
                    while ws.cell(row=row_idx, column=col_idx).value not in (None, ""):
                        row_idx += 1

                    # Bereichsname und Kostenstelle eintragen
                    ws.cell(row=row_idx, column=col_idx, value=bereichsname)
                    ws.cell(row=row_idx, column=col_idx+1, value=kostenstelle)

                    wb.save(EXCEL_PATH)
                    st.success("Bereich erfolgreich angelegt.")
            except Exception as e:
                st.error(f"Fehler beim Schreiben in Excel:\n{e}")
